from flask import Blueprint, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.colors import HexColor
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from io import BytesIO
from datetime import datetime

word_editor_bp = Blueprint('word_editor', __name__)

# Simple in-memory document storage for demo
documents = []
doc_id_counter = 1

DEFAULT_FORMATTING = {
    'fontFamily': 'Arial',
    'fontSize': 12,
    'fontColor': '#000000',
    'bold': False,
    'italic': False,
    'underline': False,
    'margins': {
        'top': 1,
        'bottom': 1,
        'left': 1,
        'right': 1
    },
    'alignment': 'left',
    'lineSpacing': 1.5,
    'pageNumbers': False
}


def merge_formatting(existing, new_values):
    formatting = existing.copy()
    if not new_values:
        return formatting

    for key, value in new_values.items():
        if key == 'margins' and isinstance(value, dict):
            merged_margins = formatting.get('margins', DEFAULT_FORMATTING['margins']).copy()
            merged_margins.update({k: float(v) for k, v in value.items() if v is not None})
            formatting['margins'] = merged_margins
        else:
            formatting[key] = value
    return formatting


def inches(value, default=1.0):
    try:
        return Inches(float(value))
    except (TypeError, ValueError):
        return Inches(default)


def parse_color(color_str):
    if not color_str:
        return None
    color = color_str.lstrip('#')
    if len(color) != 6:
        return None
    try:
        return RGBColor.from_string(color.upper())
    except ValueError:
        return None


def apply_run_formatting(run, formatting):
    font = run.font
    font.name = formatting.get('fontFamily') or DEFAULT_FORMATTING['fontFamily']

    font_size = formatting.get('fontSize')
    try:
        font.size = Pt(float(font_size))
    except (TypeError, ValueError):
        font.size = Pt(DEFAULT_FORMATTING['fontSize'])

    color = parse_color(formatting.get('fontColor'))
    if color:
        font.color.rgb = color

    run.bold = bool(formatting.get('bold'))
    run.italic = bool(formatting.get('italic'))
    run.underline = bool(formatting.get('underline'))


def apply_paragraph_formatting(paragraph, formatting):
    pf = paragraph.paragraph_format
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    alignment = formatting.get('alignment', DEFAULT_FORMATTING['alignment'])
    paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

    line_spacing = formatting.get('lineSpacing', DEFAULT_FORMATTING['lineSpacing'])
    try:
        pf.line_spacing = float(line_spacing)
    except (TypeError, ValueError):
        pf.line_spacing = DEFAULT_FORMATTING['lineSpacing']


def add_page_numbers(document):
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'PAGE'
    run._r.append(instr_text)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_end)


def build_docx_document(doc):
    document = Document()
    formatting = doc.get('formatting', DEFAULT_FORMATTING)

    # Apply section margins
    section = document.sections[0]
    margins = formatting.get('margins', DEFAULT_FORMATTING['margins'])
    section.top_margin = inches(margins.get('top', 1))
    section.bottom_margin = inches(margins.get('bottom', 1))
    section.left_margin = inches(margins.get('left', 1))
    section.right_margin = inches(margins.get('right', 1))

    # Add content
    paragraphs = [p for p in doc.get('content', '').split('\n') if p is not None]
    for para in paragraphs:
        text = para.rstrip()
        paragraph = document.add_paragraph()
        if text.strip():
            run = paragraph.add_run(text)
        else:
            run = paragraph.add_run(' ')
        apply_run_formatting(run, formatting)
        apply_paragraph_formatting(paragraph, formatting)

    if formatting.get('pageNumbers'):
        add_page_numbers(document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf_document(doc):
    formatting = doc.get('formatting', DEFAULT_FORMATTING)
    buffer = BytesIO()

    margins = formatting.get('margins', DEFAULT_FORMATTING['margins'])
    left_margin = float(margins.get('left', 1)) * 72
    right_margin = float(margins.get('right', 1)) * 72
    top_margin = float(margins.get('top', 1)) * 72
    bottom_margin = float(margins.get('bottom', 1)) * 72

    pdf_doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin
    )

    font_family_map = {
        'Arial': 'Helvetica',
        'Calibri': 'Helvetica',
        'Georgia': 'Times-Roman',
        'Times New Roman': 'Times-Roman'
    }

    base_font = font_family_map.get(formatting.get('fontFamily'), 'Helvetica')
    font_variant = base_font

    bold = formatting.get('bold')
    italic = formatting.get('italic')

    if base_font == 'Times-Roman':
        if bold and italic:
            font_variant = 'Times-BoldItalic'
        elif bold:
            font_variant = 'Times-Bold'
        elif italic:
            font_variant = 'Times-Italic'
    else:
        # Helvetica/Courier naming
        if bold and italic:
            font_variant = f'{base_font}-BoldOblique'
        elif bold:
            font_variant = f'{base_font}-Bold'
        elif italic:
            font_variant = f'{base_font}-Oblique'

    alignment_map = {
        'left': TA_LEFT,
        'center': TA_CENTER,
        'right': TA_RIGHT,
        'justify': TA_JUSTIFY
    }

    line_spacing = float(formatting.get('lineSpacing', DEFAULT_FORMATTING['lineSpacing']))
    font_size = float(formatting.get('fontSize', DEFAULT_FORMATTING['fontSize']))

    style = ParagraphStyle(
        'Custom',
        fontName=font_variant,
        fontSize=font_size,
        leading=font_size * line_spacing,
        alignment=alignment_map.get(formatting.get('alignment', 'left'), TA_LEFT),
        textColor=HexColor(formatting.get('fontColor', '#000000')),
        underline=bool(formatting.get('underline'))
    )

    elements = []
    for para in doc.get('content', '').split('\n'):
        text = para.strip()
        if not text:
            elements.append(Spacer(1, font_size * 0.6))
        else:
            elements.append(Paragraph(text.replace('\n', '<br/>'), style))
            elements.append(Spacer(1, font_size * 0.3))

    pdf_doc.build(elements or [Paragraph(' ', style)])
    buffer.seek(0)
    return buffer

@word_editor_bp.route('/upload-doc', methods=['POST'])
def upload_doc():
    try:
        print("Upload request received")
        print("Content-Type:", request.headers.get('Content-Type'))
        print("Body:", request.get_json())

        data = request.get_json()
        filename = data.get('filename', 'uploaded_document.docx') if data else 'uploaded_document.docx'
        
        # Generate realistic document content based on filename
        if 'resume' in filename.lower() or 'cv' in filename.lower():
            document_content = """RESUME

John Doe
Software Engineer
Email: john.doe@email.com
Phone: (555) 123-4567

EXPERIENCE
Senior Software Engineer at TechCorp (2020-Present)
• Developed web applications using React and Node.js
• Led a team of 5 developers
• Improved application performance by 40%

Software Engineer at StartupXYZ (2018-2020)
• Built mobile applications using React Native
• Collaborated with design team on user interfaces

EDUCATION
Bachelor of Computer Science
University of Technology (2014-2018)

SKILLS
• JavaScript, Python, Java
• React, Node.js, MongoDB
• Git, Docker, AWS

ACHIEVEMENTS
• Published 3 technical articles
• Speaker at 2 tech conferences
• AWS Certified Solutions Architect"""
        elif 'report' in filename.lower() or 'analysis' in filename.lower():
            document_content = """PROJECT REPORT

Title: Analysis of Student Performance in Online Learning

ABSTRACT
This report analyzes the impact of online learning on student performance during the academic year 2023-2024. The study examines various factors affecting student engagement and academic outcomes.

INTRODUCTION
The transition to online learning has significantly changed educational methodologies. This study examines the effectiveness of digital learning platforms and their impact on student achievement.

METHODOLOGY
Data was collected from 500 students across different departments using:
• Online surveys and questionnaires
• Academic performance records
• Engagement metrics from learning platforms
• Focus group discussions

FINDINGS
1. Student engagement decreased by 15% in online formats
2. Technical difficulties affected 30% of students
3. Self-motivated students performed better in online settings
4. Interactive content improved retention by 25%

CONCLUSION
Online learning presents both challenges and opportunities for educational institutions. Proper implementation and support systems are crucial for success.

RECOMMENDATIONS
• Improve technical infrastructure
• Provide training for both students and faculty
• Develop hybrid learning models
• Implement interactive learning tools"""
        elif 'letter' in filename.lower() or 'application' in filename.lower():
            document_content = f"""APPLICATION LETTER

[Date: {datetime.now().strftime('%Y-%m-%d')}]

Dear Hiring Manager,

I am writing to express my strong interest in the Software Engineer position at your esteemed organization. With my background in computer science and hands-on experience in web development, I am confident that I would be a valuable addition to your team.

QUALIFICATIONS
• Bachelor's degree in Computer Science
• 3+ years of experience in software development
• Proficiency in React, Node.js, and Python
• Strong problem-solving and analytical skills

ACHIEVEMENTS
In my previous role, I successfully led a team project that resulted in a 40% improvement in application performance. I have also contributed to open-source projects and maintained a strong commitment to clean, maintainable code.

I am excited about the opportunity to contribute to your organization's continued success and would welcome the chance to discuss how my skills and enthusiasm can benefit your team.

Thank you for your time and consideration.

Sincerely,
[Your Name]"""
        else:
            document_content = f"""DOCUMENT: {filename}

This is a sample document that demonstrates the Word Editor functionality.

INTRODUCTION
This document contains sample text that would normally be extracted from an uploaded .docx file. In a production environment, the system would use libraries like python-docx to extract the actual content from uploaded documents.

CONTENT SECTIONS

Section 1: Document Overview
Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris.

Section 2: Features and Capabilities
Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.

Section 3: Technical Implementation
Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo.

CONCLUSION
This document demonstrates how uploaded content would appear in the Word Editor. Users can apply various formatting options including font changes, alignment, spacing, and more.

---
Document Information:
- Filename: {filename}
- Upload Date: {datetime.now().strftime('%Y-%m-%d')}
- Processing Status: Complete"""

        global doc_id_counter
        new_doc = {
            'id': str(doc_id_counter),
            'filename': filename,
            'content': document_content,
            'uploaded_at': datetime.utcnow(),
            'formatting': DEFAULT_FORMATTING.copy()
        }

        documents.append(new_doc)
        doc_id_counter += 1

        print("Document created successfully:", new_doc['id'])

        return jsonify({
            'success': True,
            'message': 'Document uploaded successfully',
            'documentId': new_doc['id'],
            'content': document_content
        })

    except Exception as e:
        print(f"Upload document error: {e}")
        return jsonify({
            'success': False,
            'message': f'Failed to upload document: {str(e)}'
        }), 500

@word_editor_bp.route('/edit-doc', methods=['POST'])
def edit_doc():
    try:
        data = request.get_json()
        document_id = data.get('documentId')
        formatting = data.get('formatting', {})
        content = data.get('content')

        # Find document
        doc = next((d for d in documents if d['id'] == document_id), None)
        if not doc:
            return jsonify({
                'success': False,
                'message': 'Document not found'
            }), 404

        doc['formatting'] = merge_formatting(doc.get('formatting', DEFAULT_FORMATTING), formatting)
        if isinstance(content, str):
            doc['content'] = content
            print(f"Content updated. New length: {len(content)}")
            print(f"Content preview (first 200 chars): {content[:200]}")

        # In production, apply formatting using python-docx or similar
        print("Applying formatting:", doc['formatting'])
        print(f"Document state after edit - Content length: {len(doc.get('content', ''))}, Formatting keys: {list(doc.get('formatting', {}).keys())}")

        return jsonify({
            'success': True,
            'message': 'Document formatting applied successfully',
            'preview': doc['content']
        })

    except Exception as e:
        print(f"Edit document error: {e}")
        return jsonify({
            'success': False,
            'message': f'Failed to edit document: {str(e)}'
        }), 500

@word_editor_bp.route('/download-doc/<document_id>')
def download_doc(document_id):
    try:
        # Find document
        doc = next((d for d in documents if d['id'] == document_id), None)
        if not doc:
            return jsonify({
                'success': False,
                'message': 'Document not found'
            }), 404

        export_format = request.args.get('format', 'docx').lower()
        original_name = doc['filename']
        
        # Debug: Log what content and formatting we're using
        print(f"Downloading document {document_id}")
        print(f"Content length: {len(doc.get('content', ''))}")
        print(f"Content preview (first 200 chars): {doc.get('content', '')[:200]}")
        print(f"Formatting: {doc.get('formatting', {})}")

        if export_format == 'pdf':
            file_buffer = build_pdf_document(doc)
            mimetype = 'application/pdf'
            download_name = f"{original_name.rsplit('.', 1)[0]}_edited.pdf"
        else:
            file_buffer = build_docx_document(doc)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            download_name = f"{original_name.rsplit('.', 1)[0]}_edited.docx"

        return send_file(
            file_buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=download_name
        )

    except Exception as e:
        print(f"Download document error: {e}")
        return jsonify({
            'success': False,
            'message': f'Failed to download document: {str(e)}'
        }), 500
