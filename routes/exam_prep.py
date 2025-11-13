from flask import Blueprint, request, jsonify, send_file
from docx import Document
from io import BytesIO
from datetime import datetime
import random
import os
import re
import fitz  # PyMuPDF
import mammoth  # For DOCX
import requests
import json
import subprocess
import sys

exam_prep_bp = Blueprint('exam_prep', __name__)

# Simple in-memory storage for demo
uploaded_papers = []
generated_questions = []
paper_id_counter = 1
questions_id_counter = 1

# Groq API configuration
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.3-70b-versatile"

def extract_text_from_file(file_path, filename):
    """Extract text from PDF, DOCX, or TXT files"""
    ext = os.path.splitext(filename)[1].lower()
    
    print(f"Extracting text from {ext} file: {filename}")
    print(f"File path: {file_path}")
    print(f"File exists: {os.path.exists(file_path)}")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_size = os.path.getsize(file_path)
    print(f"File size: {file_size} bytes")
    
    try:
        if ext == '.pdf':
            # Use PyMuPDF (fitz) for PDF extraction
            print("=" * 60)
            print("DEBUG: Starting PDF extraction with PyMuPDF")
            print("=" * 60)
            
            # Check file first
            with open(file_path, 'rb') as f:
                first_bytes = f.read(100)
                print(f"First 100 bytes of file (hex): {first_bytes[:100].hex()}")
                print(f"First 100 bytes (ascii): {first_bytes[:100]}")
                if first_bytes.startswith(b'%PDF'):
                    print("✓ File starts with %PDF - valid PDF header")
                else:
                    print("✗ File does NOT start with %PDF - may be corrupted!")
            
            try:
                doc = fitz.open(file_path)
                print(f"✓ PDF opened successfully")
                print(f"PDF metadata:")
                print(f"  - Page count: {len(doc)}")
                print(f"  - Is encrypted: {doc.is_encrypted}")
                print(f"  - Needs password: {doc.needs_pass}")
                if doc.metadata:
                    print(f"  - Metadata: {doc.metadata}")
                
                text_parts = []
                total_chars = 0
                total_pages = len(doc)
                
                # Process ALL pages, not just first 5
                for page_num in range(total_pages):
                    try:
                        page = doc[page_num]
                        print(f"\n--- Page {page_num + 1} ---")
                        
                        # Try different extraction methods
                        text = page.get_text()
                        print(f"  get_text() result: {len(text)} chars")
                        if text:
                            print(f"  First 200 chars: {text[:200]}")
                        
                        # Try get_text("text") explicitly
                        text_alt = page.get_text("text")
                        print(f"  get_text('text') result: {len(text_alt)} chars")
                        
                        # Try get_text("dict")
                        try:
                            text_dict = page.get_text("dict")
                            blocks = text_dict.get("blocks", [])
                            print(f"  Text blocks found: {len(blocks)}")
                            for i, block in enumerate(blocks[:3]):  # First 3 blocks
                                if "lines" in block:
                                    print(f"    Block {i}: {len(block['lines'])} lines")
                        except Exception as e:
                            print(f"  Could not get text dict: {e}")
                        
                        # Check for images
                        image_list = page.get_images()
                        print(f"  Images on page: {len(image_list)}")
                        
                        if text.strip():
                            text_parts.append(text)
                            total_chars += len(text)
                        else:
                            print(f"  ⚠ No text extracted from page {page_num + 1}")
                            
                    except Exception as e:
                        print(f"  ✗ Error on page {page_num + 1}: {e}")
                        import traceback
                        traceback.print_exc()
                        continue
                
                doc.close()
                
                full_text = '\n\n'.join(text_parts)
                print("\n" + "=" * 60)
                print(f"EXTRACTION SUMMARY:")
                print(f"  Total pages in PDF: {total_pages}")
                print(f"  Pages processed: {total_pages}")
                print(f"  Pages with text: {len(text_parts)}")
                print(f"  Total characters extracted: {len(full_text)}")
                print(f"  Text preview (first 500 chars):")
                print(f"  {full_text[:500]}")
                print("=" * 60)
                
                if len(full_text.strip()) < 50:
                    print("\n⚠ WARNING: Very little text extracted!")
                    print("Possible reasons:")
                    print("  1. PDF is image-based (scanned document)")
                    print("  2. PDF is corrupted")
                    print("  3. PDF text is in images, not selectable text")
                    print("  4. PDF uses non-standard encoding")
                
                return full_text
                
            except Exception as e:
                print(f"\n✗ ERROR opening PDF: {e}")
                import traceback
                traceback.print_exc()
                raise
            
        elif ext in ['.docx', '.doc']:
            # Use mammoth for DOCX files
            with open(file_path, 'rb') as f:
                result = mammoth.extract_raw_text(f)
                return result.value
            
        elif ext == '.txt':
            # Plain text file
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
    except Exception as e:
        print(f"Error extracting text from {ext} file: {e}")
        raise

def generate_questions_with_groq(text, filename):
    """Generate questions using Groq API"""
    groq_api_key = os.getenv('GROQ_API_KEY')
    if not groq_api_key:
        raise ValueError("GROQ_API_KEY not found in environment variables")
    
    # Chunk text if too large (Groq limit is ~12k tokens, roughly 8k chars)
    # Use smart chunking to avoid splitting questions in the middle
    MAX_CHUNK_SIZE = 6000  # Smaller chunks to leave more room for complete responses
    OVERLAP_SIZE = 800  # Larger overlap to ensure questions aren't split
    chunks = []
    if len(text) > MAX_CHUNK_SIZE:
        print(f"Text too large ({len(text)} chars), chunking with smart boundaries...")
        i = 0
        while i < len(text):
            chunk_end = min(i + MAX_CHUNK_SIZE, len(text))
            chunk = text[i:chunk_end]
            
            # Try to end chunk at a good boundary to avoid splitting questions
            if chunk_end < len(text):
                # Look for question endings first (most important)
                last_question_mark = chunk.rfind('?')
                # Then look for double newlines (question separators)
                last_double_newline = chunk.rfind('\n\n')
                # Then look for single newlines
                last_newline = chunk.rfind('\n')
                # Then periods
                last_period = chunk.rfind('.')
                
                # Prefer question marks, then double newlines, then single newlines, then periods
                boundary = -1
                if last_question_mark > chunk_end - 1000:  # Within 1000 chars of end
                    boundary = last_question_mark
                elif last_double_newline > chunk_end - 800:
                    boundary = last_double_newline
                elif last_newline > chunk_end - 500:
                    boundary = last_newline
                elif last_period > chunk_end - 300:
                    boundary = last_period
                
                if boundary > chunk_end - 1000:  # If we found a good boundary
                    chunk = chunk[:boundary + 1]
                    i = i + boundary + 1 - OVERLAP_SIZE  # Overlap from previous chunk
                    print(f"  Chunk boundary at position {boundary} (question mark/newline)")
                else:
                    # No good boundary found, use standard overlap
                    i = chunk_end - OVERLAP_SIZE
            else:
                i = chunk_end
            
            if chunk.strip():
                chunks.append(chunk.strip())
        print(f"Split into {len(chunks)} chunks with smart boundaries")
    else:
        chunks = [text]
    
    all_questions = []
    
    for i, chunk in enumerate(chunks):
        print(f"Processing chunk {i + 1}/{len(chunks)}...")
        
        prompt = f"""Extract important questions from this document content and GENERATE comprehensive answers based on the document.

CRITICAL INSTRUCTIONS - READ CAREFULLY:
1. Extract COMPLETE questions with ALL parts:
   - The full question text/stem
   - ALL multiple choice options (A., B., C., D., etc.) if the question has them
   - If a question says "Which of the following" or "Select the option", you MUST include all the A, B, C, D options that follow
   - Complete sentences even if they span multiple lines
   - All question parts and sub-questions

2. For multiple choice questions, the question field should look like:
   "Which of the following statements best describes X?
   A. Option A text
   B. Option B text  
   C. Option C text
   D. Option D text"

3. NEVER truncate or cut off questions. If a question has options, include ALL of them.

4. GENERATE ANSWERS: For each question, you MUST generate a comprehensive answer based on the document content:
   - For multiple choice questions: Provide the correct option letter (A, B, C, D, etc.) AND a brief explanation of why it's correct
   - For open-ended questions: Generate a detailed answer based on the information in the document
   - For fill-in-the-blank questions: Provide the missing word/phrase and explain it
   - Answers should be accurate, detailed, and based solely on the document content provided
   - If the document doesn't contain enough information, infer a reasonable answer based on the context

5. Answer format examples:
   - Multiple choice: "C. [Explanation of why C is correct based on document]"
   - Open-ended: "[Detailed answer explaining the concept based on document content]"
   - Fill-in-the-blank: "[Missing word/phrase]. [Explanation]"

Return a JSON array of objects with these fields:
- question: string (COMPLETE question with ALL options if present - this is critical!)
- answer: string (GENERATED comprehensive answer based on document content - NEVER use "Answer not provided")
- importance: "high" | "medium" | "low"
- topic: string
- difficulty: "easy" | "medium" | "hard"
- confidence: number (0-1)

Return ONLY valid JSON array, no markdown, no explanations.

Document content:
{chunk}"""
        
        try:
            response = requests.post(
                GROQ_API_URL,
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {groq_api_key}'
                },
                json={
                    'model': GROQ_MODEL,
                    'messages': [
                    {
                        'role': 'system',
                        'content': 'You are an expert at extracting COMPLETE educational questions from documents and GENERATING comprehensive answers. CRITICAL: (1) Always include the FULL question text with ALL multiple choice options (A, B, C, D, etc.), complete sentences, and all parts of the question. Never truncate or cut off questions. (2) ALWAYS generate detailed answers based on the document content - never return "Answer not provided". For multiple choice questions, provide the correct option and explanation. For open-ended questions, provide comprehensive answers based on the document. Return only valid JSON arrays.'
                    },
                        {
                            'role': 'user',
                            'content': prompt
                        }
                    ],
                    'temperature': 0.3,
                    'max_tokens': 8000  # Increased further to ensure complete questions with all options
                },
                timeout=60
            )
            
            if not response.ok:
                error_data = response.json()
                print(f"Groq API error: {error_data}")
                raise Exception(f"Groq API error: {error_data.get('error', {}).get('message', 'Unknown error')}")
            
            data = response.json()
            response_content = data.get('choices', [{}])[0].get('message', {}).get('content', '')
            
            # Check if response was truncated
            finish_reason = data.get('choices', [{}])[0].get('finish_reason', '')
            if finish_reason == 'length':
                print(f"⚠ WARNING: Groq response was truncated (hit max_tokens limit) for chunk {i + 1}")
                print(f"Response length: {len(response_content)} characters")
                print(f"This means some questions may be incomplete. Consider using smaller chunks.")
                
                # Try to extract what we can, but warn about truncation
                # The JSON might be incomplete, so we need to handle that
            
            if not response_content:
                print(f"No content in Groq response for chunk {i + 1}")
                continue
            
            print(f"Groq response length: {len(response_content)} characters")
            print(f"Response preview (first 300 chars): {response_content[:300]}")
            
            # Parse JSON from response
            json_string = response_content.strip()
            
            # Remove markdown code blocks if present
            if json_string.startswith('```'):
                json_string = json_string.replace('```json', '').replace('```', '').strip()
            
            # Try to find JSON array in response
            json_match = None
            if '[' in json_string:
                start = json_string.index('[')
                # If response was truncated, the closing bracket might be missing
                if ']' in json_string:
                    end = json_string.rindex(']') + 1
                    json_string = json_string[start:end]
                else:
                    # Response was truncated mid-JSON, try to fix it
                    print(f"⚠ JSON response appears incomplete (missing closing bracket)")
                    # Try to find the last complete question object
                    # Look for the last complete } before the end
                    last_brace = json_string.rfind('}')
                    if last_brace > start:
                        json_string = json_string[start:last_brace + 1] + ']'
                        print(f"Attempting to fix incomplete JSON...")
                    else:
                        raise ValueError("JSON response is too incomplete to parse")
            
            try:
                questions = json.loads(json_string)
            except json.JSONDecodeError as e:
                print(f"⚠ JSON parsing error: {e}")
                print(f"JSON string length: {len(json_string)}")
                print(f"JSON string preview: {json_string[:500]}...")
                print(f"JSON string end: ...{json_string[-500:]}")
                # Try to extract partial questions if possible
                raise ValueError(f"Failed to parse JSON from Groq response: {e}")
            
            if not isinstance(questions, list):
                print(f"Groq returned non-array for chunk {i + 1}")
                continue
            
            # Normalize questions
            for q_idx, q in enumerate(questions):
                question_text = str(q.get('question', q.get('Question', ''))).strip()
                answer_text = str(q.get('answer', q.get('Answer', q.get('solution', '')))).strip()
                # If answer is empty or says "not provided", generate a placeholder
                if not answer_text or 'not provided' in answer_text.lower() or 'answer not' in answer_text.lower():
                    answer_text = 'Answer will be generated based on document content.'
                
                # Check if question seems incomplete
                if question_text:
                    # Check if question ends abruptly (no punctuation, or ends with incomplete word)
                    is_incomplete = False
                    incomplete_reason = ""
                    
                    # Check 1: Ends with incomplete words
                    if len(question_text) > 50 and not question_text[-1] in '.?!:\n' and not question_text.endswith('...'):
                        last_50 = question_text[-50:].lower()
                        if any(word in last_50 for word in ['because', 'when', 'where', 'which', 'what', 'how', 'why', 'the', 'a ', 'an ', 'and ', 'or ', 'but ']):
                            is_incomplete = True
                            incomplete_reason = f"ends with incomplete phrase: ...{question_text[-40:]}"
                    
                    # Check 2: Question mentions options but doesn't show them
                    if 'option' in question_text.lower() or 'choose' in question_text.lower() or 'select' in question_text.lower():
                        # Check if it has A., B., C., D. patterns
                        has_options = bool(re.search(r'\b[A-Z]\.\s', question_text))
                        if not has_options and len(question_text) < 200:
                            is_incomplete = True
                            incomplete_reason = "mentions options but options not included"
                    
                    # Check 3: Ends with "because" or "because ________"
                    if question_text.rstrip().endswith('because') or question_text.rstrip().endswith('because ________') or question_text.rstrip().endswith('because _____'):
                        # This might be intentional (fill-in-the-blank), but check if it's cut off
                        if not question_text.rstrip().endswith('________') and not question_text.rstrip().endswith('_____'):
                            is_incomplete = True
                            incomplete_reason = "ends with 'because' but no blank or continuation"
                    
                    if is_incomplete:
                        print(f"⚠ Question {q_idx + 1} may be incomplete: {incomplete_reason}")
                        print(f"   Full question: {question_text[:150]}...")
                
                normalized = {
                    'question': question_text,
                    'answer': answer_text,
                    'importance': str(q.get('importance', q.get('Importance', 'medium'))).lower(),
                    'topic': str(q.get('topic', q.get('Topic', 'General'))),
                    'difficulty': str(q.get('difficulty', q.get('Difficulty', 'medium'))).lower(),
                    'confidence': float(q.get('confidence', q.get('Confidence', 0.8)))
                }
                
                # Only add if question is valid
                if normalized['question'] and len(normalized['question']) > 10:
                    all_questions.append(normalized)
                    # Log question length for debugging
                    if len(normalized['question']) < 100:
                        print(f"  ✓ Question {q_idx + 1}: {len(normalized['question'])} chars - {normalized['question'][:60]}...")
                else:
                    print(f"  ✗ Skipped invalid question {q_idx + 1}: {normalized['question'][:50] if normalized['question'] else 'empty'}")
            
            print(f"✓ Got {len(questions)} questions from chunk {i + 1}")
            
        except Exception as e:
            print(f"Error processing chunk {i + 1}: {e}")
            continue
    
    return all_questions

@exam_prep_bp.route('/test', methods=['GET'])
def test():
    """Test endpoint to verify the blueprint is working"""
    return jsonify({
        'success': True,
        'message': 'Exam prep blueprint is working!',
        'routes': ['/upload-papers', '/generate-questions', '/download-questions']
    })

@exam_prep_bp.route('/upload-papers', methods=['POST'])
def upload_papers():
    global paper_id_counter
    try:
        # For demo purposes, simulate uploading previous year papers
        # In production, integrate with Google Drive API
        
        data = request.get_json()
        files = data.get('files', []) if data else []
        uploaded_files = []

        # Simulate processing files
        for file in files:
            paper = {
                'id': str(paper_id_counter),
                'filename': file.get('filename', f'previous_year_paper_{paper_id_counter}.pdf'),
                'content': "Sample question paper content",  # In production, extract from actual files
                'uploaded_at': datetime.utcnow()
            }
            
            uploaded_papers.append(paper)
            uploaded_files.append(paper)
            paper_id_counter += 1

        return jsonify({
            'success': True,
            'message': f'Successfully uploaded {len(uploaded_files)} question papers',
            'uploadedFiles': [{'id': f['id'], 'filename': f['filename']} for f in uploaded_files]
        })
    except Exception as e:
        print(f"Upload papers error: {e}")
        return jsonify({
            'success': False,
            'message': 'Failed to upload question papers'
        }), 500

@exam_prep_bp.route('/generate-questions', methods=['POST'])
def generate_questions():
    print('=' * 50)
    print('Python backend: generate_questions() called')
    print(f'Request method: {request.method}')
    print(f'Request path: {request.path}')
    print(f'Request data: {request.get_json()}')
    print('=' * 50)
    
    try:
        data = request.get_json()
        
        # Get papers from request (passed from Node.js backend)
        papers_data = data.get('papers', []) if data else []
        print(f'Received {len(papers_data)} papers from Node.js backend')
        
        # Fallback: if papers not provided, try to get from paperIds
        if not papers_data:
            paper_ids = data.get('paperIds', []) if data else []
            papers_data = [p for p in uploaded_papers if not paper_ids or p['id'] in paper_ids]
        
        print(f"Generating questions for {len(papers_data)} papers")
        
        if not papers_data:
            return jsonify({
                'success': False,
                'message': 'No papers found'
            }), 404
        
        all_questions = []
        
        # Process each paper
        for paper in papers_data:
            filename = paper.get('filename', 'unknown')
            print(f"\n=== Processing: {filename} ===")
            
            # Get file path from paper data
            file_path = paper.get('filePath')
            if not file_path or not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                continue
            
            try:
                # Step 1: Extract text from file
                print(f"Reading file: {file_path}")
                extracted_text = extract_text_from_file(file_path, filename)
                
                print(f"Extracted text length: {len(extracted_text) if extracted_text else 0}")
                if extracted_text:
                    print(f"Extracted text preview (first 500 chars): {extracted_text[:500]}")
                
                if not extracted_text or len(extracted_text.strip()) < 50:
                    print(f"ERROR: Insufficient text extracted from {filename}")
                    print(f"Text length: {len(extracted_text) if extracted_text else 0}")
                    print("This PDF might be corrupted, scanned (image-based), or encrypted.")
                    print("Trying to continue anyway...")
                    # Don't continue - try to process with what we have
                    if not extracted_text or len(extracted_text.strip()) < 10:
                        raise ValueError(f"Could not extract any readable text from {filename}. The PDF may be corrupted, scanned (image-based), or encrypted.")
                
                print(f"✓ Extracted {len(extracted_text)} characters")
                
                # Step 2: Simple cleanup
                clean_text = re.sub(r'\n{3,}', '\n\n', extracted_text)
                clean_text = re.sub(r'  +', ' ', clean_text).strip()
                print(f"Cleaned text: {len(clean_text)} characters")
                
                # Step 3: Generate questions using Groq
                print("Sending to Groq AI for question generation...")
                questions = generate_questions_with_groq(clean_text, filename)
                
                if questions:
                    all_questions.extend(questions)
                    print(f"✓ Generated {len(questions)} questions from {filename}")
                else:
                    print(f"No questions generated from {filename}")
                    
            except Exception as e:
                print(f"Error processing {filename}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        if not all_questions:
            return jsonify({
                'success': False,
                'message': 'No questions could be generated from the uploaded papers'
            }), 500
        
        # Create questions record
        global questions_id_counter
        questions_record = {
            'id': str(questions_id_counter),
            'questions': all_questions,
            'generated_at': datetime.utcnow()
        }
        
        generated_questions.append(questions_record)
        questions_id_counter += 1
        
        print(f"\n✓ Total: {len(all_questions)} questions generated")
        
        return jsonify({
            'success': True,
            'message': f'Generated {len(all_questions)} questions from {len(papers_data)} papers',
            'questions': all_questions
        })

    except Exception as e:
        print(f"Generate questions error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Failed to generate questions: {str(e)}'
        }), 500

@exam_prep_bp.route('/download-questions/<questions_id>')
def download_questions(questions_id):
    try:
        # Find questions
        questions_record = next((q for q in generated_questions if q['id'] == questions_id), None)
        if not questions_record:
            return jsonify({
                'success': False,
                'message': 'Questions not found'
            }), 404

        # Create a Word document
        doc = Document()
        doc.add_heading('Exam Preparation Questions', 0)
        doc.add_paragraph(f'Generated on: {questions_record["generated_at"].strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')

        for i, q in enumerate(questions_record['questions'], 1):
            # Add question
            doc.add_heading(f'Question {i}', level=1)
            doc.add_paragraph(q['question'])
            
            # Add importance level
            importance_para = doc.add_paragraph()
            importance_para.add_run('Importance: ').bold = True
            importance_para.add_run(q['importance'].title())
            
            # Add answer
            doc.add_heading('Answer:', level=2)
            doc.add_paragraph(q['answer'])
            doc.add_paragraph('')  # Add spacing

        # Save to BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'exam_questions_{questions_id}.docx'
        )

    except Exception as e:
        print(f"Download questions error: {e}")
        return jsonify({
            'success': False,
            'message': f'Failed to download questions: {str(e)}'
        }), 500
